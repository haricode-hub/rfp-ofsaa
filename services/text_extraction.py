"""
Text extraction utilities for different file formats.
"""

import structlog
from pathlib import Path
from typing import Optional
import re

logger = structlog.get_logger()

def extract_text_from_pdf(file_path: Path) -> str:
    """
    Extract text from PDF file.
    """
    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        logger.info("PDF text extraction completed", file_path=str(file_path), chars=len(text))
        return text.strip()

    except ImportError:
        raise Exception("pdfplumber not installed. Install with: pip install pdfplumber")
    except Exception as e:
        logger.error("PDF extraction failed", file_path=str(file_path), error=str(e))
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(file_path: Path) -> str:
    """
    Extract text from DOCX file.
    """
    try:
        from docx import Document

        doc = Document(str(file_path))
        text = ""

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"

        logger.info("DOCX text extraction completed", file_path=str(file_path), chars=len(text))
        return text.strip()

    except ImportError:
        raise Exception("python-docx not installed. Install with: pip install python-docx")
    except Exception as e:
        logger.error("DOCX extraction failed", file_path=str(file_path), error=str(e))
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

def extract_text_from_txt(file_path: Path, encoding: str = "utf-8") -> str:
    """
    Extract text from plain text file.
    """
    try:
        # Try multiple encodings if utf-8 fails
        encodings_to_try = [encoding, "utf-8", "latin-1", "cp1252"]

        for enc in encodings_to_try:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    text = f.read()
                logger.info("Text file extraction completed", file_path=str(file_path), encoding=enc, chars=len(text))
                return text
            except UnicodeDecodeError:
                continue

        raise Exception(f"Could not decode file with any of the attempted encodings: {encodings_to_try}")

    except Exception as e:
        logger.error("Text file extraction failed", file_path=str(file_path), error=str(e))
        raise Exception(f"Failed to read text file: {str(e)}")

def extract_text_from_file(file_path: Path) -> str:
    """
    Extract text from file based on extension.
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        return extract_text_from_docx(file_path)
    elif suffix in [".txt", ".md"]:
        return extract_text_from_txt(file_path)
    else:
        raise Exception(f"Unsupported file type: {suffix}. Supported types: .pdf, .docx, .txt, .md")

def validate_file_type(filename: str) -> bool:
    """
    Validate if file type is supported.
    """
    supported_extensions = {".pdf", ".docx", ".txt", ".md"}
    return Path(filename).suffix.lower() in supported_extensions