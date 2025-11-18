"""
RFP processing service combining analysis, proposal generation, and document building.
"""

import json
import re
import structlog
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, Tuple
from io import BytesIO

import httpx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml

from .config import settings
from .organization import ORG, PROPOSAL_JSON_TEMPLATE
from .rfp_models import (
    Proposal, RFPAnalysis, RFPClassification, RFPSubmissionInfo,
    UploadResponse, GenerateJSONResponse, HealthResponse, PingOllamaResponse
)
from .text_extraction import extract_text_from_file, validate_file_type
# from .ai_service import generate_completion  # Will add to ai_service

logger = structlog.get_logger()

class RFPAnalyzer:
    """Service for analyzing and classifying RFPs."""

    # Classification categories and their keywords
    CLASSIFICATION_CATALOG = {
        "Resource Augmentation": [
            "resource augmentation", "staff augmentation", "contract staffing",
            "augmentation support", "manpower supply", "resource support", "team augmentation"
        ],
        "Upgradation": [
            "upgrade", "upgradation", "modernization", "migration", "version upgrade",
            "enhancement", "technology refresh", "system upgrade"
        ],
        "Managed Service": [
            "managed service", "managed services", "operations and maintenance",
            "operation & maintenance", "support services", "outsourcing",
            "service management", "service desk", "24x7 support", "managed operations"
        ],
        "New Installation": [
            "implementation", "installation", "deploy", "deployment", "rollout",
            "greenfield", "set up", "setup", "new system", "install", "core banking system implementation"
        ]
    }

    @staticmethod
    def classify_rfp(text: str) -> RFPClassification:
        """
        Classify RFP text into categories.
        """
        if not text or not text.strip():
            return RFPClassification(
                category="Not enough information",
                confidence="low",
                matched_keywords=[]
            )

        normalized_text = text.lower()
        tallies = []

        # Score each category
        for category, keywords in RFPAnalyzer.CLASSIFICATION_CATALOG.items():
            score = 0
            hits = []

            for keyword in keywords:
                # Use word boundaries for more accurate matching
                pattern = r'\b' + re.escape(keyword.lower()) + r'\b'
                matches = re.findall(pattern, normalized_text)
                if matches:
                    score += len(matches)
                    hits.append(keyword)

            tallies.append({
                'category': category,
                'score': score,
                'hits': hits
            })

        # Sort by score (highest first)
        tallies.sort(key=lambda x: x['score'], reverse=True)
        top = tallies[0]

        if not top or top['score'] == 0:
            return RFPClassification(
                category="General Services",
                confidence="low",
                matched_keywords=[]
            )

        # Determine confidence level
        confidence = "high" if top['score'] >= 4 else "medium" if top['score'] >= 2 else "low"

        return RFPClassification(
            category=top['category'],
            confidence=confidence,
            matched_keywords=top['hits']
        )

    @staticmethod
    def extract_submission_signals(text: str) -> Dict[str, Optional[str]]:
        """
        Extract submission-related information from RFP text.
        """
        signals: Dict[str, Optional[str]] = {
            'issuance_date': None,
            'submission_deadline': None,
            'clarification_deadline': None,
            'submission_method': None,
            'contacts': None
        }

        # Define patterns for different types of information
        patterns = {
            'issuance_date': [
                r'date of issuance[:\-]\s*([^\n]+)',
                r'issuance date[:\-]\s*([^\n]+)',
                r'issue date[:\-]\s*([^\n]+)'
            ],
            'submission_deadline': [
                r'submission deadline[:\-]\s*([^\n]+)',
                r'proposal submission deadline[:\-]\s*([^\n]+)',
                r'submission date[:\-]\s*([^\n]+)',
                r'closing date[:\-]\s*([^\n]+)',
                r'last date for submission[:\-]\s*([^\n]+)',
                r'deadline[:\-]\s*([^\n]+)'
            ],
            'clarification_deadline': [
                r'clarification deadline[:\-]\s*([^\n]+)',
                r'clarifications deadline[:\-]\s*([^\n]+)',
                r'questions deadline[:\-]\s*([^\n]+)',
                r'query deadline[:\-]\s*([^\n]+)',
                r'last date for clarifications[:\-]\s*([^\n]+)'
            ],
            'submission_method': [
                r'submission method[:\-]\s*([^\n]+)',
                r'mode of submission[:\-]\s*([^\n]+)',
                r'submission process[:\-]\s*([^\n]+)',
                r'submission email[:\-]\s*([^\n]+)',
                r'submission address[:\-]\s*([^\n]+)'
            ],
            'contacts': [
                r'contact person[:\-]\s*([^\n]+)',
                r'contact[:\-]\s*([^\n]+)',
                r'email[:\-]\s*([^\n]+)',
                r'point of contact[:\-]\s*([^\n]+)'
            ]
        }

        for field, field_patterns in patterns.items():
            for pattern in field_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match and match.group(1):
                    signals[field] = match.group(1).strip()
                    break

        return signals

    @staticmethod
    def create_submission_info(signals: Dict[str, Optional[str]]) -> RFPSubmissionInfo:
        """
        Create RFPSubmissionInfo from extracted signals.
        """
        return RFPSubmissionInfo(
            issuance_date=signals.get('issuance_date'),
            submission_deadline=signals.get('submission_deadline'),
            clarification_deadline=signals.get('clarification_deadline'),
            submission_method=signals.get('submission_method'),
            contacts=signals.get('contacts')
        )

    @staticmethod
    def generate_summary_snippet(text: str, max_length: int = 420) -> str:
        """
        Generate a summary snippet from RFP text.
        """
        if not text:
            return ""

        import re
        cleaned = re.sub(r'\s+', ' ', text).strip()
        if len(cleaned) <= max_length:
            return cleaned

        return cleaned[:max_length] + "…"

    @classmethod
    async def analyze_rfp_text(cls, text: str, existing_classification: Optional[RFPClassification] = None) -> RFPAnalysis:
        """
        Perform complete RFP analysis.
        """
        try:
            # Get or create classification
            base_classification = existing_classification or cls.classify_rfp(text)

            # Extract submission signals
            submission_signals = cls.extract_submission_signals(text)
            submission_info = cls.create_submission_info(submission_signals)

            # Generate summary
            summary = cls.generate_summary_snippet(text)

            # For now, return basic analysis structure
            # In a full implementation, this would use AI to extract detailed information
            analysis = RFPAnalysis(
                rfp_type=base_classification,
                summary=summary,
                issuing_organization=None,  # Would be extracted by AI
                scope=None,  # Would be extracted by AI
                functional_requirements=[],  # Would be extracted by AI
                technical_requirements=[],  # Would be extracted by AI
                services=[],  # Would be extracted by AI
                submission=submission_info,
                evaluation_focus=[],  # Would be extracted by AI
                optional_components=[],  # Would be extracted by AI
                risks=[]  # Would be extracted by AI
            )

            logger.info("RFP analysis completed", category=base_classification.category)
            return analysis

        except Exception as e:
            logger.error("RFP analysis failed", error=str(e))
            # Return minimal analysis on failure
            fallback_classification = existing_classification or RFPClassification(
                category="General Services",
                confidence="low",
                matched_keywords=[]
            )

            return RFPAnalysis(
                rfp_type=fallback_classification,
                summary=cls.generate_summary_snippet(text),
                submission=cls.create_submission_info({}),
                functional_requirements=[],
                technical_requirements=[],
                services=[],
                evaluation_focus=[],
                optional_components=[],
                risks=[]
            )

class ProposalGenerator:
    """Service for generating proposals using AI."""

    def __init__(self):
        """Initialize the proposal generator."""
        self.client = httpx.AsyncClient(timeout=settings.ollama_timeout)

    async def generate_proposal_json(self, rfp_text: str, meta: Optional[Dict[str, Any]] = None) -> Proposal:
        """
        Generate proposal JSON from RFP text using AI.
        """
        try:
            # Prepare the prompt
            system_prompt = f"You are a senior proposal writer for {ORG.name}. Produce clean JSON only."
            user_prompt = f"""From the following RFP text, extract/compose a JMR proposal structure. Return **ONLY** valid JSON matching:
{PROPOSAL_JSON_TEMPLATE}

RFP TEXT:

{rfp_text}

Keep the narrative polished and concise. Ensure payment_milestones percentages sum to 100 and amount equals percent of the overall grand total (rounded to two decimals). If some fields are missing, infer reasonable defaults for a mid-size banking/fintech implementation. Avoid hallucinations about specific vendor SKUs unless clearly in the RFP."""

            full_prompt = f"{system_prompt}\n\n{user_prompt}"

            # Use ai_service for generation if available, else direct Ollama
            # Direct Ollama call for now
            payload = {
                "model": settings.ollama_model,
                "stream": False,
                "prompt": full_prompt
            }
            async with self.client as client:
                response = await client.post(
                    f"{settings.ollama_base_url}/api/generate",
                    json=payload,
                    timeout=settings.ollama_timeout
                )

                if response.status_code != 200:
                    raise Exception(f"Ollama API error: {response.status_code} - {response.text}")

                response_data = response.json()
                raw_response = response_data.get("response", "").strip()

            if not raw_response:
                raise Exception("Empty response from AI")

            # Extract JSON from response
            json_match = raw_response
            if not json_match.startswith('{'):
                json_pattern = r'\{[\s\S]*\}'
                match = re.search(json_pattern, json_match)
                if match:
                    json_match = match.group(0)
                else:
                    raise Exception("No JSON found in AI response")

            # Parse JSON
            proposal_data = json.loads(json_match)

            # Create Proposal object
            proposal = Proposal(**proposal_data)

            # Apply metadata overrides
            if meta:
                if 'client_name' in meta:
                    proposal.client_name = meta['client_name']
                if 'project_title' in meta:
                    proposal.project_title = meta['project_title']

            # Normalize and validate proposal
            proposal = self._normalize_proposal(proposal)

            logger.info("Proposal generated successfully", client=proposal.client_name)
            return proposal

        except Exception as e:
            logger.error("Proposal generation failed", error=str(e))
            # Fall back to local generation
            return self._create_fallback_proposal(rfp_text, meta or {})

    def _normalize_proposal(self, proposal: Proposal) -> Proposal:
        """
        Normalize and validate proposal data.
        """
        # Ensure arrays are properly formatted
        proposal.scope_of_work = [item for item in (proposal.scope_of_work or []) if item and item.strip()]
        proposal.deliverables = [item for item in (proposal.deliverables or []) if item and item.strip()]
        proposal.prerequisites = [item for item in (proposal.prerequisites or []) if item and item.strip()]
        proposal.scope_exclusions = [item for item in (proposal.scope_exclusions or []) if item and item.strip()]
        proposal.assumptions = [item for item in (proposal.assumptions or []) if item and item.strip()]
        proposal.customer_obligations = [item for item in (proposal.customer_obligations or []) if item and item.strip()]
        proposal.acceptance_criteria = [item for item in (proposal.acceptance_criteria or []) if item and item.strip()]
        proposal.payment_terms_details = [item for item in (proposal.payment_terms_details or []) if item and item.strip()]

        # Normalize commercials
        if proposal.commercials:
            comm = proposal.commercials
            comm.currency = comm.currency or ORG.currency
            comm.discount_percent = max(0, min(100, comm.discount_percent or 0))
            comm.tax_percent = comm.tax_percent or ORG.default_tax

            # Calculate totals
            subtotal = sum((item.qty or 0) * (item.rate or 0) for item in (comm.line_items or []))
            discount_amount = subtotal * (comm.discount_percent / 100)
            taxable = subtotal - discount_amount
            tax_amount = taxable * (comm.tax_percent / 100)
            grand_total = taxable + tax_amount

            # Normalize payment milestones
            if comm.payment_milestones:
                total_percent = sum(milestone.percent for milestone in comm.payment_milestones)
                if total_percent != 100:
                    # Adjust percentages to sum to 100
                    factor = 100 / total_percent if total_percent > 0 else 1
                    for milestone in comm.payment_milestones:
                        milestone.percent = round(milestone.percent * factor, 2)
                        milestone.amount = round(grand_total * (milestone.percent / 100), 2)

        # Set defaults for missing fields
        if not proposal.validity:
            proposal.validity = "30 days from issue date"

        return proposal

    def _create_fallback_proposal(self, rfp_text: str, meta: Dict[str, Any]) -> Proposal:
        """
        Create a fallback proposal when AI generation fails.
        """
        logger.info("Creating fallback proposal")

        # Extract snippet from RFP text
        snippet = re.sub(r'\s+', ' ', rfp_text).strip()[:400] if rfp_text else ""

        # Default line items
        line_items = [
            {
                "name": "Oracle Cloud Professional Services",
                "unit": "Engagement",
                "qty": 1,
                "rate": 120000
            }
        ]

        # Calculate totals
        subtotal = sum(item["qty"] * item["rate"] for item in line_items)
        tax_percent = ORG.default_tax
        tax_amount = subtotal * (tax_percent / 100)
        grand_total = subtotal + tax_amount

        return Proposal(
            client_name=meta.get('client_name', 'Client'),
            project_title=meta.get('project_title', 'Oracle Implementation Proposal (Auto-generated)'),
            executive_summary=f"This Oracle-focused proposal was auto-generated from the provided RFP. It outlines how JMR Infotech accelerates value on Oracle platforms. RFP excerpt: {snippet}",
            scope_of_work=[
                "Oracle Cloud discovery and fit-gap assessment across priority business processes",
                "Solution design, configuration, and controlled extensions leveraging Oracle best practices",
                "Data migration, integration enablement, and readiness for Oracle Cloud environments",
                "End-to-end testing, cutover orchestration, and hypercare support for Oracle workloads"
            ],
            deliverables=[
                "Oracle solution blueprint and configuration workbooks",
                "Provisioned Oracle Cloud test and production environments",
                "Integration and migration runbooks with validation sign-offs",
                "Knowledge transfer and enablement sessions for Oracle operations teams"
            ],
            prerequisites=[
                "Customer confirms Oracle Cloud subscriptions and grants environment access ahead of discovery",
                "Required Oracle licenses, CSI numbers, and approvals are available before configuration begins",
                "Client subject-matter experts for each Oracle process area are identified for workshops"
            ],
            scope_exclusions=[
                "Procurement of new Oracle or third-party licenses beyond agreed scope",
                "Custom development outside approved Oracle extension frameworks",
                "Legacy system remediation that is unrelated to the Oracle programme"
            ],
            assumptions=[
                "Client provides timely access to Oracle Cloud consoles, integration endpoints, and decision-makers",
                "Standard Oracle adapters and accelerators are sufficient unless otherwise documented",
                "Third-party vendors collaborate on Oracle integration touchpoints as needed"
            ],
            customer_obligations=[
                "Nominate Oracle product owners and process leads for each functional tower",
                "Provide connectivity, VPN, and Oracle Service Request access for project resources",
                "Review and sign off Oracle design, build, and test deliverables within agreed timelines"
            ],
            timeline=[
                {
                    "phase": "Oracle Mobilization",
                    "duration": "3 weeks",
                    "milestones": ["Mobilization complete", "Oracle discovery workshops finished"]
                },
                {
                    "phase": "Solution Build",
                    "duration": "4 months",
                    "milestones": ["Configuration baselined", "Integration cycle complete"]
                },
                {
                    "phase": "Transition & Hypercare",
                    "duration": "4 weeks",
                    "milestones": ["Production cutover executed", "Oracle hypercare exit"]
                }
            ],
            resource_plan=[
                {
                    "role": "Oracle Program Manager",
                    "count": 1,
                    "mode": "onsite"
                },
                {
                    "role": "Oracle Solution Architect",
                    "count": 1,
                    "mode": "offshore"
                },
                {
                    "role": "Oracle Technical Consultant",
                    "count": 2,
                    "mode": "offshore"
                },
                {
                    "role": "Oracle Integration Specialist",
                    "count": 1,
                    "mode": "offshore"
                }
            ],
            commercials=CommercialInfo(
                currency=ORG.currency,
                line_items=[CommercialLineItem(**item) for item in line_items],
                discount_percent=0,
                tax_percent=tax_percent,
                payment_terms_summary="Oracle delivery services billed against milestone completions; invoices due net 30 days.",
                out_of_pocket_expenses=[
                    "Economy airfare and local transit in line with Oracle delivery governance",
                    "Hotel accommodation aligned with client travel policy during Oracle workshops",
                    "Visa and work permit costs for Oracle delivery resources, if applicable"
                ],
                payment_milestones=[
                    PaymentMilestone(
                        description="Oracle mobilization and environment readiness",
                        percent=30,
                        amount=round(grand_total * 0.30, 2)
                    ),
                    PaymentMilestone(
                        description="Solution build and conference room pilot sign-off",
                        percent=40,
                        amount=round(grand_total * 0.40, 2)
                    ),
                    PaymentMilestone(
                        description="Production go-live and Oracle hypercare closure",
                        percent=30,
                        amount=round(grand_total * 0.30, 2)
                    )
                ]
            ),
            payment_terms_details=[
                "Invoices payable within 30 calendar days from submission; Oracle Partner Network terms apply.",
                "Out-of-pocket expenses billed at actuals upon submission of receipts and client approval.",
                "All commercials exclude applicable taxes, duties, and Oracle subscription fees unless stated otherwise."
            ],
            acceptance_criteria=[
                "Configured Oracle Cloud modules approved by business process owners",
                "Successful completion of SIT and UAT cycles with client sign-off",
                "Operational handover of Oracle environments and knowledge transfer completed"
            ],
            validity="30 days from issue date for Oracle commercial terms."
        )

class DocumentBuilder:
    """Service for building DOCX documents from proposal data."""

    def __init__(self):
        """Initialize the document builder."""
        try:
            import docx
            self.docx_available = True
        except ImportError:
            logger.warning("python-docx not installed. DOCX generation will not be available.")
            self.docx_available = False

    async def build_docx_from_proposal(self, proposal: Proposal, today: Optional[str] = None) -> bytes:
        """
        Build a DOCX document from proposal data.
        """
        if not self.docx_available:
            raise Exception("python-docx library not available. Install with: pip install python-docx")

        if not today:
            today = datetime.now().strftime('%Y-%m-%d')

        try:
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import nsdecls, qn
            from docx.oxml import parse_xml

            # Create document
            doc = Document()

            # Set up styles
            self._setup_styles(doc)

            # Add cover page
            self._add_cover_page(doc, proposal, today)

            # Add page break
            doc.add_page_break()

            # Add content sections
            self._add_executive_summary(doc, proposal)
            self._add_scope_of_work(doc, proposal)
            self._add_deliverables(doc, proposal)
            self._add_prerequisites(doc, proposal)
            self._add_scope_exclusions(doc, proposal)
            self._add_assumptions(doc, proposal)
            self._add_customer_obligations(doc, proposal)
            self._add_resource_plan(doc, proposal)
            self._add_timeline(doc, proposal)
            self._add_commercials(doc, proposal)
            self._add_payment_terms(doc, proposal)
            self._add_acceptance_criteria(doc, proposal)
            self._add_validity(doc, proposal)

            # Set page margins and headers/footers
            self._setup_page_layout(doc, proposal, today)

            # Save to bytes
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            docx_content = buffer.getvalue()
            logger.info("DOCX document generated successfully", size=len(docx_content))

            return docx_content

        except Exception as e:
            logger.error("DOCX generation failed", error=str(e))
            raise Exception(f"Failed to generate DOCX: {str(e)}")

    def _setup_styles(self, doc):
        """Set up document styles."""
        from docx.shared import Pt
        from docx.enum.style import WD_STYLE_TYPE

        # Title style
        title_style = doc.styles.add_style('JMRTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.name = 'Calibri'

        # Heading styles
        for i in range(1, 4):
            heading_style = doc.styles.add_style(f'JMRHeading{i}', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.size = Pt(18 - i * 2)
            heading_style.font.bold = True
            heading_style.font.name = 'Calibri'

        # Body styles
        body_style = doc.styles.add_style('JMRBody', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.size = Pt(11)
        body_style.font.name = 'Calibri'

        bullet_style = doc.styles.add_style('JMRBullet', WD_STYLE_TYPE.PARAGRAPH)
        bullet_style.font.size = Pt(11)
        bullet_style.font.name = 'Calibri'

    def _add_cover_page(self, doc, proposal: Proposal, today: str):
        """Add cover page to document."""
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Add some spacing
        for _ in range(8):
            doc.add_paragraph()

        # Title
        title = doc.add_paragraph(proposal.project_title or 'Proposal', style='JMRTitle')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = doc.add_paragraph(ORG.tagline, style='JMRBody')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing
        for _ in range(4):
            doc.add_paragraph()

        # Client info
        client_info = doc.add_paragraph(style='JMRBody')
        client_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        client_info.add_run(f"Prepared for: {proposal.client_name or 'Client'}\n")
        client_info.add_run(f"Prepared by: {ORG.name}\n")
        client_info.add_run(f"Date: {today}\n")
        client_info.add_run(f"{ORG.website} | {ORG.email} | {ORG.phone}\n")
        client_info.add_run(ORG.address)

    def _add_executive_summary(self, doc, proposal: Proposal):
        """Add executive summary section."""
        doc.add_paragraph('Executive Summary', style='JMRHeading1')
        summary = proposal.executive_summary or 'Executive summary to be provided.'
        doc.add_paragraph(summary, style='JMRBody')

    def _add_scope_of_work(self, doc, proposal: Proposal):
        """Add scope of work section."""
        doc.add_paragraph('Scope of Work', style='JMRHeading1')
        scope_items = proposal.scope_of_work or ['Scope details to be confirmed.']
        for item in scope_items:
            p = doc.add_paragraph(item, style='JMRBullet')
            p.paragraph_format.left_indent = Inches(0.25)

    def _add_deliverables(self, doc, proposal: Proposal):
        """Add deliverables section."""
        doc.add_paragraph('Deliverables', style='JMRHeading1')
        deliverables = proposal.deliverables or ['Deliverables to be confirmed.']
        for item in deliverables:
            p = doc.add_paragraph(item, style='JMRBullet')
            p.paragraph_format.left_indent = Inches(0.25)

    def _add_prerequisites(self, doc, proposal: Proposal):
        """Add prerequisites section."""
        doc.add_paragraph('Prerequisites', style='JMRHeading1')
        prerequisites = proposal.prerequisites or ['Prerequisites to be confirmed.']
        for item in prerequisites:
            p = doc.add_paragraph(item, style='JMRBullet')
            p.paragraph_format.left_indent = Inches(0.25)

    def _add_scope_exclusions(self, doc, proposal: Proposal):
        """Add scope exclusions section."""
        doc.add_paragraph('Scope Exclusions', style='JMRHeading1')
        exclusions = proposal.scope_exclusions or ['Exclusions to be confirmed.']
        for item in exclusions:
            p = doc.add_paragraph(item, style='JMRBullet')
            p.paragraph_format.left_indent = Inches(0.25)

    def _add_assumptions(self, doc, proposal: Proposal):
        """Add assumptions section."""
        doc.add_paragraph('Assumptions', style='JMRHeading1')
        assumptions = proposal.assumptions or ['Assumptions to be confirmed.']
        for item in assumptions:
            p = doc.add_paragraph(item, style='JMRBullet')
            p.paragraph_format.left_indent = Inches(0.25)

    def _add_customer_obligations(self, doc, proposal: Proposal):
        """Add customer obligations section."""
        doc.add_paragraph('Customer Obligations', style='JMRHeading1')
        obligations = proposal.customer_obligations or ['Customer obligations to be confirmed.']
        for item in obligations:
            p = doc.add_paragraph(item, style='JMRBullet')
            p.paragraph_format.left_indent = Inches(0.25)

    def _add_resource_plan(self, doc, proposal: Proposal):
        """Add resource plan section with table."""
        doc.add_paragraph('Resource Plan', style='JMRHeading1')

        if not proposal.resource_plan:
            doc.add_paragraph('Resource plan to be confirmed.', style='JMRBody')
            return

        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Role'
        header_cells[1].text = 'Count'
        header_cells[2].text = 'Mode'

        # Data rows
        for resource in proposal.resource_plan:
            row_cells = table.add_row().cells
            row_cells[0].text = resource.role or 'TBD'
            row_cells[1].text = str(resource.count or 1)
            row_cells[2].text = resource.mode or 'TBD'

    def _add_timeline(self, doc, proposal: Proposal):
        """Add project timeline section with table."""
        doc.add_paragraph('Project Timeline', style='JMRHeading1')

        if not proposal.timeline:
            doc.add_paragraph('Timeline to be confirmed.', style='JMRBody')
            return

        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Phase'
        header_cells[1].text = 'Duration'
        header_cells[2].text = 'Milestones'

        # Data rows
        for phase in proposal.timeline:
            row_cells = table.add_row().cells
            row_cells[0].text = phase.phase or 'TBD'
            row_cells[1].text = phase.duration or 'TBD'

            # Milestones as bullet points
            milestones_text = '\n'.join(f'• {m}' for m in (phase.milestones or ['TBD']))
            row_cells[2].text = milestones_text

    def _add_commercials(self, doc, proposal: Proposal):
        """Add commercials section with detailed breakdown."""
        doc.add_paragraph('Commercials', style='JMRHeading1')

        if not proposal.commercials:
            doc.add_paragraph('Commercial details to be confirmed.', style='JMRBody')
            return

        comm = proposal.commercials
        currency = comm.currency or ORG.currency

        # Line items table
        if comm.line_items:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Description'
            header_cells[1].text = 'Unit'
            header_cells[2].text = 'Qty'
            header_cells[3].text = 'Rate'
            header_cells[4].text = 'Amount'

            # Data rows
            for item in comm.line_items:
                row_cells = table.add_row().cells
                qty = item.qty or 0
                rate = item.rate or 0
                amount = qty * rate

                row_cells[0].text = item.name or 'Service'
                row_cells[1].text = item.unit or 'LOE'
                row_cells[2].text = str(qty)
                row_cells[3].text = f"{currency} {rate:,.2f}"
                row_cells[4].text = f"{currency} {amount:,.2f}"

        # Calculate totals
        subtotal = sum((item.qty or 0) * (item.rate or 0) for item in (comm.line_items or []))
        discount = subtotal * ((comm.discount_percent or 0) / 100)
        taxable = subtotal - discount
        tax_amount = taxable * ((comm.tax_percent or ORG.default_tax) / 100)
        grand_total = taxable + tax_amount

        # Summary
        doc.add_paragraph(f"Subtotal: {currency} {subtotal:,.2f}", style='JMRBody')
        if discount > 0:
            doc.add_paragraph(f"Discount ({comm.discount_percent:.1f}%): -{currency} {discount:,.2f}", style='JMRBody')
        doc.add_paragraph(f"Taxable Amount: {currency} {taxable:,.2f}", style='JMRBody')
        doc.add_paragraph(f"Tax ({comm.tax_percent or ORG.default_tax:.1f}%): {currency} {tax_amount:,.2f}", style='JMRBody')
        doc.add_paragraph(f"Grand Total: {currency} {grand_total:,.2f}", style='JMRBody')

    def _add_payment_terms(self, doc, proposal: Proposal):
        """Add payment terms section."""
        doc.add_paragraph('Payment Terms & Conditions', style='JMRHeading1')

        payment_details = proposal.payment_terms_details or [
            'Invoices payable within 30 calendar days from submission.',
            'Out-of-pocket expenses billed at actuals with supporting receipts.'
        ]

        for term in payment_details:
            p = doc.add_paragraph(term, style='JMRBullet')
            p.paragraph_format.left_indent = Inches(0.25)

        # Payment milestones
        if proposal.commercials and proposal.commercials.payment_milestones:
            doc.add_paragraph('Payment Milestones', style='JMRHeading2')

            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Description'
            header_cells[1].text = '% Payment'
            header_cells[2].text = 'Amount'

            # Data rows
            currency = proposal.commercials.currency or ORG.currency
            for milestone in proposal.commercials.payment_milestones:
                row_cells = table.add_row().cells
                row_cells[0].text = milestone.description or 'Milestone'
                row_cells[1].text = f"{milestone.percent:.1f}%"
                row_cells[2].text = f"{currency} {milestone.amount:,.2f}"

    def _add_acceptance_criteria(self, doc, proposal: Proposal):
        """Add acceptance criteria section."""
        doc.add_paragraph('Acceptance Criteria', style='JMRHeading1')
        criteria = proposal.acceptance_criteria or ['Acceptance criteria to be confirmed.']
        for item in criteria:
            p = doc.add_paragraph(item, style='JMRBullet')
            p.paragraph_format.left_indent = Inches(0.25)

    def _add_validity(self, doc, proposal: Proposal):
        """Add proposal validity section."""
        doc.add_paragraph('Proposal Validity', style='JMRHeading1')
        validity = proposal.validity or f'Commercial terms are valid for 30 days from {datetime.now().strftime("%B %d, %Y")}.'
        doc.add_paragraph(validity, style='JMRBody')

    def _setup_page_layout(self, doc, proposal: Proposal, today: str):
        """Set up page layout, headers, and footers."""
        from docx.shared import Inches

        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

            # Add header
            header = section.header
            header_paragraph = header.paragraphs[0]
            header_paragraph.text = f"{ORG.name} | {ORG.tagline}"
            header_paragraph.style = 'Header'

            # Add footer
            footer = section.footer
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.text = f"{ORG.website} | {ORG.email} | {ORG.phone} | Page "
            footer_paragraph.style = 'Footer'

            # Add page number (simplified)
            footer_paragraph.add_run(" of ")  # Basic, can enhance if needed

# Global instances
rfp_analyzer = RFPAnalyzer()
proposal_generator = ProposalGenerator()
document_builder = DocumentBuilder()

# Core functions for rfp.py
async def analyze_rfp(file_path: Path) -> UploadResponse:
    """
    Analyze RFP file: extract text, classify, and analyze.
    """
    if not validate_file_type(file_path.name):
        raise Exception(f"Unsupported file type: {file_path.suffix}")

    text = extract_text_from_file(file_path)
    if len(text) > settings.analysis_max_chars:
        text = text[:settings.analysis_max_chars]

    classification = rfp_analyzer.classify_rfp(text)
    analysis = await rfp_analyzer.analyze_rfp_text(text, classification)

    preview = text[:2000] if len(text) > 2000 else text

    return UploadResponse(
        ok=True,
        chars=len(text),
        preview=preview,
        rfp_text=text,
        classification=classification,
        analysis=analysis
    )

async def generate_proposal(analysis_data: Dict[str, Any]) -> GenerateJSONResponse:
    """
    Generate proposal JSON from analysis.
    """
    rfp_text = analysis_data.get('rfp_text', '')
    meta = analysis_data.get('meta', {})

    proposal = await proposal_generator.generate_proposal_json(rfp_text, meta)

    return GenerateJSONResponse(
        ok=True,
        proposal=proposal
    )

async def build_docx(proposal_data: Proposal) -> bytes:
    """
    Build DOCX from proposal.
    """
    return await document_builder.build_docx_from_proposal(proposal_data)

async def health_check() -> HealthResponse:
    """
    Health check for RFP service.
    """
    return HealthResponse(
        status="healthy",
        service="rfp-service"
    )

async def ping_ollama() -> PingOllamaResponse:
    """
    Check Ollama connectivity.
    """
    try:
        client = httpx.AsyncClient(timeout=5)
        async with client as c:
            response = await c.get(f"{settings.ollama_base_url}/api/tags")
        return PingOllamaResponse(ok=True)
    except Exception as e:
        return PingOllamaResponse(ok=False, error=str(e))