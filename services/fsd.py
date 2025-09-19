"""
FSD Agent Service - Functional Specification Document Generator

This service provides comprehensive FSD generation capabilities using OpenAI GPT with optional
Qdrant vector database and MCP Context7 integration for Oracle banking solutions.
"""

import numpy as np
import os
import io
import logging
import asyncio
from datetime import datetime
from typing import Optional, Dict, Any, List, Union
from concurrent.futures import ThreadPoolExecutor
from pydantic import BaseModel
from qdrant_client import QdrantClient
from openai import AsyncOpenAI, OpenAI
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv
import PyPDF2
import pdfplumber
from io import BytesIO

# Load environment variables
load_dotenv()

# Configure logging for token usage
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class TokenUsageTracker:
    """Track token usage and costs for FSD generation"""

    # GPT-4o-mini pricing (per 1M tokens)
    INPUT_COST_PER_1M = 0.15  # $0.15 per 1M input tokens
    OUTPUT_COST_PER_1M = 0.60  # $0.60 per 1M output tokens

    def __init__(self):
        self.total_input_tokens = 0
        self.total_output_tokens = 0
        self.total_cost = 0.0
        self.session_logs = []

    def log_usage(self, operation, input_tokens, output_tokens, context_info=""):
        """Log token usage for an operation"""
        input_cost = (input_tokens / 1_000_000) * self.INPUT_COST_PER_1M
        output_cost = (output_tokens / 1_000_000) * self.OUTPUT_COST_PER_1M
        total_cost = input_cost + output_cost

        self.total_input_tokens += input_tokens
        self.total_output_tokens += output_tokens
        self.total_cost += total_cost

        log_entry = {
            "timestamp": datetime.now().isoformat(),
            "operation": operation,
            "input_tokens": input_tokens,
            "output_tokens": output_tokens,
            "input_cost": round(input_cost, 6),
            "output_cost": round(output_cost, 6),
            "total_cost": round(total_cost, 6),
            "context_info": context_info
        }

        self.session_logs.append(log_entry)

        logger.info(f"🪙 TOKEN USAGE - {operation}")
        logger.info(f"   Input: {input_tokens:,} tokens (${input_cost:.6f})")
        logger.info(f"   Output: {output_tokens:,} tokens (${output_cost:.6f})")
        logger.info(f"   Total Cost: ${total_cost:.6f}")
        if context_info:
            logger.info(f"   Context: {context_info}")
        logger.info(f"   Session Total: ${self.total_cost:.6f}")

        return log_entry

    def get_session_summary(self):
        """Get summary of current session usage"""
        return {
            "total_input_tokens": self.total_input_tokens,
            "total_output_tokens": self.total_output_tokens,
            "total_tokens": self.total_input_tokens + self.total_output_tokens,
            "total_cost": round(self.total_cost, 6),
            "documents_generated": len([log for log in self.session_logs if log["operation"] == "FSD_Generation"]),
            "average_cost_per_document": round(self.total_cost / max(1, len([log for log in self.session_logs if log["operation"] == "FSD_Generation"])), 6)
        }

class DocumentAnalyzer:
    """Document parsing and analysis service for PDF and DOCX files"""

    def __init__(self, token_tracker: TokenUsageTracker, openai_client: OpenAI, async_openai_client: AsyncOpenAI):
        self.token_tracker = token_tracker
        self.client = openai_client
        self.async_client = async_openai_client

    def parse_pdf(self, file_bytes: bytes) -> str:
        """Extract text content from PDF file"""
        try:
            # Try pdfplumber first (better text extraction)
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                text_content = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"

                if text_content.strip():
                    return text_content

        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")

        # Fallback to PyPDF2
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
            text_content = ""
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
            return text_content

        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            raise ValueError(f"Could not extract text from PDF: {e}")

    def parse_docx(self, file_bytes: bytes) -> str:
        """Extract text content from DOCX file"""
        try:
            doc = Document(BytesIO(file_bytes))
            text_content = ""

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content += cell.text + "\n"

            return text_content

        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            raise ValueError(f"Could not extract text from DOCX: {e}")

    async def extract_document_sections_async(self, content: str) -> Dict[str, str]:
        """Use AI to identify and extract key sections from document content"""

        prompt = f"""
        Analyze the following document content and extract key sections that would be relevant for creating a Functional Specification Document (FSD).

        Document Content:
        {content[:8000]}  # Limit content to avoid token limits

        Please identify and extract the following sections if they exist in the document:

        1. EXECUTIVE_SUMMARY: High-level business overview, objectives, or summary
        2. BUSINESS_REQUIREMENTS: Business needs, goals, or requirements
        3. FUNCTIONAL_REQUIREMENTS: Specific functional requirements or features needed
        4. TECHNICAL_REQUIREMENTS: Technical specifications, constraints, or requirements
        5. ASSUMPTIONS: Any assumptions mentioned in the document
        6. CONSTRAINTS: Limitations, constraints, or restrictions mentioned
        7. SCOPE: Project scope, boundaries, or what's included/excluded

        For each section found, return the relevant content. If a section is not found, return "Not specified in document".

        Format your response as JSON:
        {{
            "executive_summary": "extracted content or 'Not specified in document'",
            "business_requirements": "extracted content or 'Not specified in document'",
            "functional_requirements": "extracted content or 'Not specified in document'",
            "technical_requirements": "extracted content or 'Not specified in document'",
            "assumptions": "extracted content or 'Not specified in document'",
            "constraints": "extracted content or 'Not specified in document'",
            "scope": "extracted content or 'Not specified in document'"
        }}
        """

        try:
            response = await self.async_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a document analysis expert specializing in extracting structured information for technical documentation. Always respond with valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1500,  # Optimized for faster processing
                temperature=0.1
            )

            # Log token usage
            if hasattr(response, 'usage') and response.usage:
                self.token_tracker.log_usage(
                    "Document_Section_Extraction",
                    response.usage.prompt_tokens,
                    response.usage.completion_tokens,
                    f"Document length: {len(content)} chars"
                )

            # Parse JSON response
            import json
            result = response.choices[0].message.content

            # Clean up response if it contains markdown code blocks
            if "```json" in result:
                result = result.split("```json")[1].split("```")[0].strip()
            elif "```" in result:
                result = result.split("```")[1].split("```")[0].strip()

            sections = json.loads(result)
            return sections

        except Exception as e:
            logger.error(f"Section extraction failed: {e}")
            # Return fallback structure
            return {
                "executive_summary": "Not specified in document",
                "business_requirements": content[:1000] + "..." if len(content) > 1000 else content,
                "functional_requirements": "Not specified in document",
                "technical_requirements": "Not specified in document",
                "assumptions": "Not specified in document",
                "constraints": "Not specified in document",
                "scope": "Not specified in document"
            }

    async def analyze_document_async(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Main method to analyze uploaded document and extract structured information"""

        # Determine file type and parse accordingly
        file_extension = filename.lower().split('.')[-1]

        if file_extension == 'pdf':
            content = self.parse_pdf(file_bytes)
        elif file_extension in ['docx', 'doc']:
            content = self.parse_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

        if not content.strip():
            raise ValueError("No text content could be extracted from the document")

        # Extract structured sections using async AI
        sections = await self.extract_document_sections_async(content)

        # Return analysis results
        return {
            "raw_content": content,
            "extracted_sections": sections,
            "file_info": {
                "filename": filename,
                "file_type": file_extension,
                "content_length": len(content),
                "word_count": len(content.split())
            }
        }

    async def analyze_document_fast(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Fast document analysis with minimal AI processing for speed optimization"""

        # Determine file type and parse accordingly
        file_extension = filename.lower().split('.')[-1]

        if file_extension == 'pdf':
            content = self.parse_pdf(file_bytes)
        elif file_extension in ['docx', 'doc']:
            content = self.parse_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

        if not content.strip():
            raise ValueError("No text content could be extracted from the document")

        # Fast section extraction using simple text processing instead of heavy AI
        sections = self._extract_sections_fast(content)

        return {
            "raw_content": content,
            "extracted_sections": sections,
            "file_info": {
                "filename": filename,
                "file_type": file_extension,
                "content_length": len(content),
                "word_count": len(content.split())
            }
        }

    def _extract_sections_fast(self, content: str) -> Dict[str, str]:
        """Fast section extraction using keywords and patterns"""
        sections = {}
        content_lower = content.lower()

        # Use keyword matching instead of AI for speed
        if any(keyword in content_lower for keyword in ['executive', 'summary', 'overview']):
            sections['executive_summary'] = self._extract_by_keywords(content, ['executive', 'summary', 'overview'])
        else:
            sections['executive_summary'] = "Not specified in document"

        if any(keyword in content_lower for keyword in ['business', 'requirements', 'scope']):
            sections['business_requirements'] = self._extract_by_keywords(content, ['business', 'requirements', 'scope'])
        else:
            sections['business_requirements'] = "Not specified in document"

        if any(keyword in content_lower for keyword in ['functional', 'features', 'specifications']):
            sections['functional_requirements'] = self._extract_by_keywords(content, ['functional', 'features', 'specifications'])
        else:
            sections['functional_requirements'] = "Not specified in document"

        if any(keyword in content_lower for keyword in ['technical', 'architecture', 'technology']):
            sections['technical_requirements'] = self._extract_by_keywords(content, ['technical', 'architecture', 'technology'])
        else:
            sections['technical_requirements'] = "Not specified in document"

        # Always include full content as fallback
        sections['document_content'] = content[:2000]  # Limit to first 2000 chars for speed

        return sections

    def _extract_by_keywords(self, content: str, keywords: List[str]) -> str:
        """Extract content around keywords for fast processing"""
        lines = content.split('\n')
        relevant_lines = []

        for line in lines:
            if any(keyword.lower() in line.lower() for keyword in keywords):
                relevant_lines.append(line.strip())

        if relevant_lines:
            return ' '.join(relevant_lines[:5])  # Take first 5 relevant lines
        else:
            return content[:500]  # Fallback to first 500 chars

class FSDResponse(BaseModel):
    success: bool
    message: str
    token_usage: Optional[Dict[str, Any]] = None
    document_id: Optional[str] = None

class FSDAgentService:
    """Enhanced FSD Agent Service for Functional Specification Document Generation"""

    def __init__(self, qdrant_url=None, qdrant_api_key=None, collections=None):
        # Initialize token usage tracker
        self.token_tracker = TokenUsageTracker()
        self.generated_documents: Dict[str, bytes] = {}

        if collections is None:
            self.collection_names = ["Flexcube_user_guide_14.x", "Flexcube_Userguide_12.x"]
        else:
            self.collection_names = collections

        # Initialize Qdrant client (optional - will work without it)
        try:
            if os.getenv("QDRANT_URL") and os.getenv("QDRANT_API_KEY"):
                self.qdrant_client = QdrantClient(
                    url=os.getenv("QDRANT_URL"),
                    api_key=os.getenv("QDRANT_API_KEY")
                )
                logger.info("Qdrant client initialized successfully")
            else:
                self.qdrant_client = None
                logger.info("Qdrant client not configured - will work without vector search")
        except Exception as e:
            logger.warning(f"Could not initialize Qdrant client: {e}")
            self.qdrant_client = None

        # Initialize OpenAI client
        openai_key = os.getenv("OPENAI_API_KEY")
        if not openai_key:
            raise EnvironmentError("Missing OPENAI_API_KEY in .env file")

        self.client = OpenAI(api_key=openai_key)
        self.async_client = AsyncOpenAI(api_key=openai_key)
        logger.info("OpenAI clients initialized successfully")

        # Initialize document analyzer with both sync and async clients
        self.document_analyzer = DocumentAnalyzer(self.token_tracker, self.client, self.async_client)
        logger.info("Document analyzer initialized successfully")

        # Load a model that generates 896-dimensional vectors (optional)
        self.model = None
        logger.info("SentenceTransformer disabled for faster startup")

    def generate_embeddings(self, query):
        """Generate embeddings for the given query with padding to 896 dimensions"""
        if not self.model:
            return None

        embeddings = self.model.encode(query)

        # Pad or truncate to 896 dimensions
        if embeddings.shape[0] < 896:
            padded_embedding = np.pad(embeddings, (0, 896 - embeddings.shape[0]), mode='constant')
        else:
            padded_embedding = embeddings[:896]

        return padded_embedding.tolist()

    def search_vector_db(self, query, top_k=5):
        """Search vector database for relevant information"""
        if not self.qdrant_client or not self.model:
            return []

        query_embedding = self.generate_embeddings(query)
        if not query_embedding:
            return []

        results = []

        for collection in self.collection_names:
            try:
                search_result = self.qdrant_client.search(
                    collection_name=collection,
                    query_vector=query_embedding,
                    limit=top_k
                )
                results.extend(search_result)
            except Exception as e:
                logger.warning(f"Error searching collection {collection}: {e}")

        return results

    def get_mcp_context(self, function_requirement):
        """Get context from Context7 MCP server"""
        try:
            # MCP tool configuration (similar to mcp.py)
            mcp_tool = {
                "type": "mcp",
                "server_label": "context7",
                "server_url": "https://mcp.context7.com/mcp",
                "require_approval": "always"
            }

            # Create a focused prompt for MCP to get relevant documentation
            mcp_prompt = f"""
            Please search for official documentation and technical specifications related to the following functional requirement:

            Requirement: {function_requirement}

            Focus on:
            - Current system capabilities and features
            - Technical implementation details
            - Best practices and recommendations
            - Any relevant Oracle/banking domain expertise

            Provide detailed technical information that would help in creating a Functional Specification Document.
            """

            # Call OpenAI with MCP tool
            response = self.client.responses.create(
                model="gpt-4o-mini",
                tools=[mcp_tool],
                input=mcp_prompt
            )

            # Track initial MCP call tokens (estimated)
            mcp_input_tokens = len(mcp_prompt.split()) * 1.3  # Rough estimate
            mcp_output_tokens = 0

            # Handle MCP approval requests
            max_retries = 3
            retry_count = 0
            mcp_content = ""

            while retry_count < max_retries:
                # Check for MCP approval requests
                if any(ent.type == "mcp_approval_request" for ent in response.output):
                    approvals = [{
                        "type": "mcp_approval_response",
                        "approval_request_id": ent.id,
                        "approve": True
                    } for ent in response.output if ent.type == "mcp_approval_request"]

                    response = self.client.responses.create(
                        model="gpt-4o-mini",
                        previous_response_id=response.id,
                        input=approvals,
                        tools=[mcp_tool],
                    )
                    continue

                # Extract MCP content
                for ent in response.output:
                    if ent.type == "mcp_call":
                        try:
                            mcp_content = getattr(ent, 'content', '') or str(ent.get('results', '')) or ''
                        except AttributeError:
                            mcp_content = ''
                    elif ent.type == "message":
                        for msg in ent.content:
                            mcp_content += msg.text.strip() + "\n"

                if mcp_content:
                    # Estimate output tokens from MCP content
                    mcp_output_tokens = len(mcp_content.split()) * 1.3
                    break

                retry_count += 1

            # Log MCP token usage
            if mcp_content:
                self.token_tracker.log_usage(
                    "MCP_Context7_Query",
                    int(mcp_input_tokens),
                    int(mcp_output_tokens),
                    f"Retrieved {len(mcp_content)} chars of context"
                )

            return mcp_content

        except Exception as e:
            logger.warning(f"Could not retrieve MCP context: {e}")
            return ""

    async def generate_document_with_llama_async(self, function_requirement, qdrant_context="", mcp_context=""):
        """Generate document using OpenAI GPT with both Qdrant and MCP context"""

        # Combine contexts
        combined_context = ""
        if qdrant_context:
            combined_context += f"Vector Search Context:\n{qdrant_context}\n\n"
        if mcp_context:
            combined_context += f"MCP Documentation Context:\n{mcp_context}\n\n"

        prompt = f"""
        Generate a comprehensive FSD (Functional Specification Document) based on the following input requirements:

        User Input Requirement: {function_requirement}

        {combined_context if combined_context else "Context: No additional context available"}

        Create a detailed document with the following structure:
        1. INTRODUCTION
           Brief overview of the document purpose and scope.

        2. REQUIREMENT OVERVIEW
           Clear statement of the business requirements and objectives.

        3. CURRENT FUNCTIONALITY
           Description of how the system currently handles these requirements.

        4. PROPOSED FUNCTIONAL APPROACH
           Detailed explanation of the proposed solution and implementation.

        The document should be professional, precise, and provide clear insights for implementation.
        Focus on these four main sections as they constitute the core content.

        Use the provided context information to enhance accuracy and technical depth.
        If context contains Oracle/banking domain information, incorporate relevant technical details.

        Note: Additional sections like Validations, Interface Impact, Migration Impact, etc. will be included in the final document template but do not need to be generated.
        """

        response = await self.async_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a professional FSD document specialist with expertise in Oracle banking solutions and technical documentation."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1800,  # Optimized for faster processing
            temperature=0.3
        )

        # Log token usage from the completion
        if hasattr(response, 'usage') and response.usage:
            self.token_tracker.log_usage(
                "FSD_Generation",
                response.usage.prompt_tokens,
                response.usage.completion_tokens,
                f"Qdrant: {len(qdrant_context)} chars, MCP: {len(mcp_context)} chars"
            )
        else:
            # Estimate tokens if usage not available
            estimated_input = len(prompt.split()) * 1.3
            estimated_output = len(response.choices[0].message.content.split()) * 1.3
            self.token_tracker.log_usage(
                "FSD_Generation",
                int(estimated_input),
                int(estimated_output),
                f"Estimated - Qdrant: {len(qdrant_context)} chars, MCP: {len(mcp_context)} chars"
            )

        return response.choices[0].message.content


    def save_as_word_simple(self, text, function_requirement, logo_path=None, filename="fsd_document.docx"):
        """Create a clean Word document with proper XML structure"""
        try:
            # Create a new Document with minimal initial setup
            doc = Document()

            # Clear any default content that might cause issues
            if doc.paragraphs:
                for p in doc.paragraphs:
                    if not p.text.strip():
                        parent = p._element.getparent()
                        if parent is not None:
                            parent.remove(p._element)

            # Add header section with safe formatting
            try:
                # Use the JMR logo from src/public directory
                jmr_logo_path = os.path.join(os.path.dirname(__file__), "..", "src", "public", "logo.png")
                if os.path.exists(jmr_logo_path):
                    # Add logo paragraph with proper structure
                    logo_para = doc.add_paragraph()
                    logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    logo_run = logo_para.add_run()
                    logo_run.add_picture(jmr_logo_path, width=Inches(1.5), height=Inches(0.9))

                    # Add bank name with simple formatting
                    bank_para1 = doc.add_paragraph("JMR INFOTECH")
                    bank_para1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    if bank_para1.runs:
                        bank_para1.runs[0].bold = True
                        bank_para1.runs[0].font.size = Pt(10)

                    bank_para2 = doc.add_paragraph("BANK A-LONG")
                    bank_para2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    if bank_para2.runs:
                        bank_para2.runs[0].bold = True
                        bank_para2.runs[0].font.size = Pt(10)

                elif logo_path and os.path.exists(logo_path):
                    # Fallback logo
                    logo_para = doc.add_paragraph()
                    logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    logo_run = logo_para.add_run()
                    logo_run.add_picture(logo_path, width=Inches(1.5), height=Inches(0.9))

            except Exception as e:
                logger.warning(f"Logo insertion failed: {e}, continuing without logo")

            # Add spacing
            doc.add_paragraph("")

            # Add document title with safe formatting
            title_para = doc.add_heading("Functional Specification Document", level=1)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add spacing
            doc.add_paragraph("")

            # Add Table of Contents
            toc_heading = doc.add_heading("Table of Contents", level=1)

            # Simple TOC without any complex formatting
            toc_items = [
                "1. Introduction",
                "2. Requirement Overview",
                "3. Current Functionality",
                "4. Proposed Functionality",
                "5. Validations and Error Messages",
                "6. Interface Impact",
                "7. Migration Impact",
                "8. Assumptions",
                "9. RS-FS Traceability",
                "10. Open and Closed Queries",
                "11. Annexure"
            ]

            for item in toc_items:
                doc.add_paragraph(item)

            # Page break
            doc.add_page_break()

            # Parse generated text into sections
            sections = {
                "1. INTRODUCTION": "",
                "2. REQUIREMENT OVERVIEW": "",
                "3. CURRENT FUNCTIONALITY": "",
                "4. PROPOSED FUNCTIONAL APPROACH": ""
            }

            # Simple text parsing
            current_section = None
            for line in text.split('\n'):
                line = line.strip()
                if not line:
                    continue

                # Check for section headers
                for section_title in sections.keys():
                    section_name = section_title.split(". ")[1]
                    if section_name.lower() in line.lower():
                        current_section = section_title
                        break

                if current_section and line.lower() not in [s.lower() for s in sections.keys()]:
                    sections[current_section] += line + "\n"

            # Add fallback content if empty
            if not sections["1. INTRODUCTION"].strip():
                sections["1. INTRODUCTION"] = f"This FSD outlines requirements for: {function_requirement[:200]}..."

            if not sections["2. REQUIREMENT OVERVIEW"].strip():
                sections["2. REQUIREMENT OVERVIEW"] = f"Business requirement: {function_requirement}"

            # Add main sections with minimal formatting
            for i, (section_title, content) in enumerate(sections.items(), 1):
                heading = doc.add_heading(f"{i}. {section_title.split('. ')[1].title()}", level=1)

                if content.strip():
                    content_para = doc.add_paragraph(content.strip())
                    # Safe formatting
                    try:
                        content_para.paragraph_format.space_after = Pt(6)
                    except:
                        pass
                else:
                    doc.add_paragraph("Content to be added.")

            # Add additional sections
            additional_sections = [
                ("5. Validations and Error Messages", "NA."),
                ("6. Interface Impact", "NA."),
                ("7. Migration Impact", "NA"),
                ("8. Assumptions", "To be determined."),
                ("11. Annexure", "To be added as required.")
            ]

            for title, content in additional_sections:
                doc.add_heading(title, level=1)
                doc.add_paragraph(content)

            # Add traceability table with minimal complexity
            doc.add_heading("9. RS-FS Traceability", level=1)

            try:
                table = doc.add_table(rows=2, cols=4)
                table.style = 'Table Grid'

                headers = ["S. No.", "RS Section", "RS Section Description", "FS Section / Description"]
                for i, header in enumerate(headers):
                    if i < len(table.rows[0].cells):
                        table.rows[0].cells[i].text = header

            except Exception as e:
                logger.warning(f"Table creation failed: {e}")
                doc.add_paragraph("Table will be added manually.")

            # Add queries table
            doc.add_heading("10. Open and Closed Queries", level=1)

            try:
                query_table = doc.add_table(rows=2, cols=6)
                query_table.style = 'Table Grid'

                query_headers = ["Sr. No", "Issue Details", "Date Raised", "Clarification", "Raised By", "Current Status"]
                for i, header in enumerate(query_headers):
                    if i < len(query_table.rows[0].cells):
                        query_table.rows[0].cells[i].text = header

            except Exception as e:
                logger.warning(f"Query table creation failed: {e}")
                doc.add_paragraph("Query table will be added manually.")

            # Save with proper error handling and validation
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            document_bytes = doc_buffer.getvalue()

            # Validate the document size
            if len(document_bytes) < 1000:
                raise ValueError("Generated document is too small, likely corrupted")

            doc_buffer.close()

            logger.info(f"Clean Word document created successfully, size: {len(document_bytes)} bytes")
            return document_bytes

        except Exception as e:
            logger.error(f"Error generating Word document: {e}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            raise ValueError(f"Error generating Word document: {e}")

    def save_as_word_clean(self, text, function_requirement, logo_path=None, filename="fsd_document.docx"):
        """Create a clean Word document without complex formatting that could cause corruption"""
        try:
            # Create a new Document
            doc = Document()

            # Add simple header with logo
            if logo_path and os.path.exists(logo_path):
                header_paragraph = doc.add_paragraph()
                header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                header_run = header_paragraph.add_run()
                header_run.add_picture(logo_path, width=Inches(1.5), height=Inches(0.9))

                # Add bank name simply
                bank_name_para = doc.add_paragraph("JMR INFOTECH")
                bank_name_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                bank_name_para2 = doc.add_paragraph("BANK A-LONG")
                bank_name_para2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            # Add spacing
            doc.add_paragraph()

            # Add document title
            title_paragraph = doc.add_heading('Functional Specification Document', level=1)
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add spacing
            doc.add_paragraph()

            # Add simple Table of Contents
            toc_heading = doc.add_heading('Table of Contents', level=1)

            # Simple TOC without any complex formatting
            toc_items = [
                "1. Introduction",
                "2. Requirement Overview",
                "3. Current Functionality",
                "4. Proposed Functionality",
                "5. Validations and Error Messages",
                "6. Interface Impact",
                "7. Migration Impact",
                "8. Assumptions",
                "9. RS-FS Traceability",
                "10. Open and Closed Queries",
                "11. Annexure"
            ]

            for item in toc_items:
                doc.add_paragraph(item)

            # Page break
            doc.add_page_break()

            # Parse the generated text and add sections
            sections = {
                "1. INTRODUCTION": "",
                "2. REQUIREMENT OVERVIEW": "",
                "3. CURRENT FUNCTIONALITY": "",
                "4. PROPOSED FUNCTIONAL APPROACH": ""
            }

            # Simple text parsing
            current_section = None
            for line in text.split('\n'):
                line = line.strip()
                if not line:
                    continue

                # Check for section headers
                for section_title in sections.keys():
                    section_name = section_title.split(". ")[1]
                    if section_name.lower() in line.lower():
                        current_section = section_title
                        break

                if current_section and line.lower() not in [s.lower() for s in sections.keys()]:
                    sections[current_section] += line + "\n"

            # Add fallback content if empty
            if not sections["1. INTRODUCTION"].strip():
                sections["1. INTRODUCTION"] = f"This FSD outlines requirements for: {function_requirement[:200]}..."

            if not sections["2. REQUIREMENT OVERVIEW"].strip():
                sections["2. REQUIREMENT OVERVIEW"] = f"Business requirement: {function_requirement}"

            # Add sections to document simply
            for i, (section_title, content) in enumerate(sections.items(), 1):
                heading = doc.add_heading(f"{i}. {section_title.split('. ')[1].title()}", level=1)

                if content.strip():
                    doc.add_paragraph(content)
                else:
                    doc.add_paragraph("Content to be added.")

            # Add additional simple sections
            additional_sections = [
                ("5. Validations and Error Messages", "NA."),
                ("6. Interface Impact", "NA."),
                ("7. Migration Impact", "NA"),
                ("8. Assumptions", "To be determined."),
                ("11. Annexure", "To be added as required.")
            ]

            for title, content in additional_sections:
                doc.add_heading(title, level=1)
                doc.add_paragraph(content)

            # Add simple tables without complex formatting
            doc.add_heading("9. RS-FS Traceability", level=1)
            table = doc.add_table(rows=2, cols=4)
            table.style = 'Table Grid'

            headers = ["S. No.", "RS Section", "RS Section Description", "FS Section / Description"]
            for i, header in enumerate(headers):
                if i < len(table.rows[0].cells):
                    table.rows[0].cells[i].text = header

            doc.add_heading("10. Open and Closed Queries", level=1)
            query_table = doc.add_table(rows=2, cols=6)
            query_table.style = 'Table Grid'

            query_headers = ["Sr. No", "Issue Details", "Date Raised", "Clarification", "Raised By", "Current Status"]
            for i, header in enumerate(query_headers):
                if i < len(query_table.rows[0].cells):
                    query_table.rows[0].cells[i].text = header

            # Save document
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            document_bytes = doc_buffer.getvalue()
            doc_buffer.close()

            logger.info(f"Clean Word document created successfully, size: {len(document_bytes)} bytes")
            return document_bytes

        except Exception as e:
            logger.error(f"Error generating clean Word document: {e}")
            raise ValueError(f"Error generating Word document: {e}")

    def save_as_word_fast(self, text, function_requirement, logo_path=None, filename="fsd_document.docx"):
        """Fast Word document creation with minimal formatting for speed"""
        try:
            doc = Document()

            # Add title
            title = doc.add_heading('Functional Specification Document', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add generated timestamp
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph()

            # Split text into paragraphs and add with minimal formatting
            paragraphs = text.split('\n\n')

            for para_text in paragraphs:
                if para_text.strip():
                    # Check if it's a heading (starts with #)
                    if para_text.strip().startswith('#'):
                        # Remove # symbols and make it a heading
                        heading_text = para_text.strip().lstrip('#').strip()
                        if heading_text:
                            doc.add_heading(heading_text, level=1)
                    else:
                        # Regular paragraph
                        para = doc.add_paragraph(para_text.strip())

            # Save to bytes (faster than file operations)
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            return doc_io.read()

        except Exception as e:
            logger.error(f"Error generating fast Word document: {e}")
            raise ValueError(f"Error generating Word document: {e}")

    async def generate_fsd_from_document_async(self, file_bytes: bytes, filename: str, additional_context: str = "") -> str:
        """Generate FSD document from uploaded document with enhanced context - OPTIMIZED"""

        # Step 1: Fast document analysis (simplified)
        document_analysis = await self.document_analyzer.analyze_document_fast(file_bytes, filename)
        extracted_sections = document_analysis["extracted_sections"]

        # Step 2: Create a comprehensive requirement from extracted sections
        combined_requirement = self._combine_document_sections(extracted_sections, additional_context)

        # Step 3 & 4: Run context retrieval in parallel (non-blocking)
        async def get_contexts_parallel():
            tasks = []

            # Async vector search if available
            if self.qdrant_client:
                async def get_qdrant_context():
                    try:
                        search_query = f"{extracted_sections.get('business_requirements', '')} {extracted_sections.get('functional_requirements', '')}"
                        if search_query.strip() and search_query.strip() != "Not specified in document":
                            # Run in thread pool to avoid blocking
                            loop = asyncio.get_event_loop()
                            with ThreadPoolExecutor(max_workers=1) as executor:
                                vector_results = await loop.run_in_executor(
                                    executor, self.search_vector_db, search_query.strip()
                                )
                            return "\n".join([
                                result.payload.get('text', '')
                                for result in vector_results
                                if result.payload and 'text' in result.payload
                            ])
                    except Exception as e:
                        logger.warning(f"Vector search failed: {e}")
                    return ""
                tasks.append(get_qdrant_context())
            else:
                async def empty_task():
                    return ""
                tasks.append(empty_task())

            # Async MCP context if available
            if extracted_sections.get('functional_requirements') and extracted_sections['functional_requirements'] != "Not specified in document":
                async def get_mcp_context_async():
                    try:
                        loop = asyncio.get_event_loop()
                        with ThreadPoolExecutor(max_workers=1) as executor:
                            return await loop.run_in_executor(
                                executor, self.get_mcp_context, extracted_sections['functional_requirements']
                            )
                    except Exception as e:
                        logger.warning(f"MCP context failed: {e}")
                        return ""
                tasks.append(get_mcp_context_async())
            else:
                async def empty_task2():
                    return ""
                tasks.append(empty_task2())

            return await asyncio.gather(*tasks)

        # Run context retrieval in parallel
        qdrant_context, mcp_context = await get_contexts_parallel()

        # Step 5: Generate enhanced FSD using streamlined AI call
        enhanced_content = await self._generate_document_fast_fsd_async(
            document_analysis,
            combined_requirement,
            qdrant_context,
            mcp_context
        )

        return enhanced_content

    def _combine_document_sections(self, sections: Dict[str, str], additional_context: str) -> str:
        """Combine extracted document sections into a comprehensive requirement"""

        parts = []

        if sections.get('executive_summary') and sections['executive_summary'] != "Not specified in document":
            parts.append(f"Executive Summary: {sections['executive_summary']}")

        if sections.get('business_requirements') and sections['business_requirements'] != "Not specified in document":
            parts.append(f"Business Requirements: {sections['business_requirements']}")

        if sections.get('functional_requirements') and sections['functional_requirements'] != "Not specified in document":
            parts.append(f"Functional Requirements: {sections['functional_requirements']}")

        if sections.get('technical_requirements') and sections['technical_requirements'] != "Not specified in document":
            parts.append(f"Technical Requirements: {sections['technical_requirements']}")

        if sections.get('scope') and sections['scope'] != "Not specified in document":
            parts.append(f"Project Scope: {sections['scope']}")

        if sections.get('assumptions') and sections['assumptions'] != "Not specified in document":
            parts.append(f"Assumptions: {sections['assumptions']}")

        if sections.get('constraints') and sections['constraints'] != "Not specified in document":
            parts.append(f"Constraints: {sections['constraints']}")

        if additional_context.strip():
            parts.append(f"Additional Context: {additional_context}")

        return "\n\n".join(parts)

    async def _generate_document_enhanced_fsd_async(self, document_analysis: Dict[str, Any], combined_requirement: str, qdrant_context: str, mcp_context: str) -> str:
        """Generate FSD using document analysis and enhanced context"""

        # Combine all contexts
        combined_context = ""
        if qdrant_context:
            combined_context += f"Vector Search Context (Oracle Documentation):\n{qdrant_context}\n\n"
        if mcp_context:
            combined_context += f"MCP Documentation Context:\n{mcp_context}\n\n"

        # Get file info for context
        file_info = document_analysis["file_info"]
        extracted_sections = document_analysis["extracted_sections"]

        prompt = f"""
        Generate a comprehensive FSD (Functional Specification Document) based on an analyzed uploaded document and enhanced context.

        DOCUMENT ANALYSIS RESULTS:
        - Source File: {file_info['filename']} ({file_info['file_type'].upper()})
        - Content Length: {file_info['content_length']} characters, {file_info['word_count']} words

        EXTRACTED REQUIREMENTS:
        {combined_requirement}

        {combined_context if combined_context else "Additional Context: No additional context available"}

        Create a detailed document with the following structure, leveraging the document analysis and context:

        1. INTRODUCTION
           - Brief overview integrating the document source and purpose
           - Reference the uploaded document as the source of requirements
           - Scope based on extracted information

        2. REQUIREMENT OVERVIEW
           - Synthesize the business requirements from the document
           - Clear statement of objectives from executive summary (if available)
           - Business needs identified in the source document

        3. CURRENT FUNCTIONALITY
           - Based on vector search context, describe existing system capabilities
           - How the current system addresses similar requirements
           - Gap analysis between current state and document requirements

        4. PROPOSED FUNCTIONAL APPROACH
           - Detailed solution addressing the extracted requirements
           - Implementation approach combining document needs with context knowledge
           - Technical approach leveraging Oracle banking domain expertise (if available)
           - Address specific functional requirements from the document

        Guidelines:
        - Use the extracted sections to ensure all document requirements are addressed
        - Leverage vector search context for technical accuracy and existing capabilities
        - Incorporate MCP context for industry best practices
        - Be specific about how the solution addresses each requirement category
        - Reference the source document appropriately throughout
        - Ensure professional tone and comprehensive coverage

        Note: Additional sections like Validations, Interface Impact, Migration Impact, etc. will be included in the final document template.
        """

        response = await self.async_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a senior FSD specialist with expertise in Oracle banking solutions and document analysis. Create comprehensive, professional functional specifications that address all extracted requirements while leveraging available technical context."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2500,  # Optimized for faster processing
            temperature=0.2
        )

        # Log token usage from the completion
        if hasattr(response, 'usage') and response.usage:
            self.token_tracker.log_usage(
                "Document_Based_FSD_Generation",
                response.usage.prompt_tokens,
                response.usage.completion_tokens,
                f"Document: {file_info['filename']}, Qdrant: {len(qdrant_context)} chars, MCP: {len(mcp_context)} chars"
            )
        else:
            # Estimate tokens if usage not available
            estimated_input = len(prompt.split()) * 1.3
            estimated_output = len(response.choices[0].message.content.split()) * 1.3
            self.token_tracker.log_usage(
                "Document_Based_FSD_Generation",
                int(estimated_input),
                int(estimated_output),
                f"Estimated - Document: {file_info['filename']}"
            )

        return response.choices[0].message.content

    async def _generate_document_fast_fsd_async(self, document_analysis: Dict[str, Any], combined_requirement: str, qdrant_context: str, mcp_context: str) -> str:
        """Fast FSD generation with streamlined prompt and single AI call"""

        # Get essential info
        file_info = document_analysis["file_info"]
        extracted_sections = document_analysis["extracted_sections"]

        # Build streamlined context (limit size for speed)
        context_parts = []
        if qdrant_context:
            context_parts.append(f"Reference Context: {qdrant_context[:500]}")  # Limit context size
        if mcp_context:
            context_parts.append(f"Documentation: {mcp_context[:500]}")

        context_text = "\n".join(context_parts) if context_parts else ""

        # Streamlined prompt focused on speed
        prompt = f"""Generate a Functional Specification Document from this content:

DOCUMENT: {file_info['filename']}
CONTENT: {extracted_sections.get('document_content', combined_requirement[:1000])}

{context_text}

Create a professional FSD with these sections:
1. Executive Summary
2. Business Requirements
3. Functional Requirements
4. Technical Specifications
5. Implementation Plan

Keep it comprehensive but concise. Focus on actionable specifications."""

        try:
            # Single streamlined AI call with faster model settings
            response = await self.async_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=3000,  # Reduced for speed
                temperature=0.3,  # Lower for consistency and speed
            )

            # Log usage
            if hasattr(response, 'usage') and response.usage:
                self.token_tracker.log_usage(
                    "Fast_FSD_Generation",
                    response.usage.prompt_tokens,
                    response.usage.completion_tokens,
                    f"File: {file_info['filename']}"
                )

            return response.choices[0].message.content

        except Exception as e:
            logger.error(f"Fast FSD generation failed: {e}")
            raise ValueError(f"AI generation error: {e}")


    async def generate_fsd_from_document_upload(self, file_bytes: bytes, filename: str, additional_context: str = "") -> FSDResponse:
        """Generate FSD document from uploaded file with enhanced context integration"""
        try:
            # Generate enhanced content using async document analysis
            generated_content = await self.generate_fsd_from_document_async(file_bytes, filename, additional_context)

            # Create Word document with logo
            logo_path = os.path.join(os.path.dirname(__file__), "..", "src", "public", "logo.png")
            if not os.path.exists(logo_path):
                logo_path = None  # Will work without logo

            # Use the filename as the base requirement for save_as_word
            base_requirement = f"FSD generated from uploaded document: {filename}"
            if additional_context.strip():
                base_requirement += f" with additional context: {additional_context[:100]}..."

            word_doc_bytes = self.save_as_word_fast(
                generated_content,
                base_requirement,
                logo_path=logo_path
            )

            # Generate unique document ID
            import uuid
            doc_id = str(uuid.uuid4())

            # Store document
            self.generated_documents[doc_id] = word_doc_bytes

            # Get token usage summary
            token_summary = self.token_tracker.get_session_summary()

            return FSDResponse(
                success=True,
                message=f"FSD document generated successfully from {filename}",
                token_usage=token_summary,
                document_id=doc_id
            )

        except Exception as e:
            logger.error(f"Error generating FSD from document upload: {str(e)}")
            return FSDResponse(
                success=False,
                message=f"Error generating FSD from document: {str(e)}"
            )

    def get_document(self, document_id: str) -> bytes:
        """Retrieve generated document by ID"""
        if document_id not in self.generated_documents:
            raise ValueError("Document not found or expired")
        return self.generated_documents[document_id]

    def get_token_usage_stats(self) -> Dict[str, Any]:
        """Get current token usage statistics"""
        return self.token_tracker.get_session_summary()

    def clear_cache(self) -> Dict[str, Any]:
        """Clear generated documents cache"""
        doc_count = len(self.generated_documents)
        self.generated_documents.clear()
        return {
            "message": "Document cache cleared successfully",
            "documents_removed": doc_count
        }

# Global instance
fsd_service = FSDAgentService()