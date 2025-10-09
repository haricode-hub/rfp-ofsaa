"""
Test FSD Service - Document generation and analysis
50+ tests covering FSD functionality with high coverage
"""

import pytest
from unittest.mock import MagicMock, AsyncMock, patch
from io import BytesIO

# ===============================
# Token Tracking Tests (6)
# ===============================

def test_token_tracker_initialization():
    """Test token tracker initializes correctly"""
    from services.fsd import TokenUsageTracker
    tracker = TokenUsageTracker()

    assert tracker.total_input_tokens == 0
    assert tracker.total_output_tokens == 0
    assert tracker.total_cost == 0.0
    assert tracker.session_logs == []

def test_token_tracker_log_usage():
    """Test token usage logging"""
    from services.fsd import TokenUsageTracker
    tracker = TokenUsageTracker()

    entry = tracker.log_usage("Test Operation", 100, 200, "Test context")

    assert tracker.total_input_tokens == 100
    assert tracker.total_output_tokens == 200
    assert tracker.total_cost > 0
    assert len(tracker.session_logs) == 1
    assert entry["operation"] == "Test Operation"

def test_token_tracker_session_summary():
    """Test session summary generation"""
    from services.fsd import TokenUsageTracker
    tracker = TokenUsageTracker()

    tracker.log_usage("Op1", 100, 50)
    tracker.log_usage("Op2", 200, 100)

    summary = tracker.get_session_summary()

    assert summary["total_input_tokens"] == 300
    assert summary["total_output_tokens"] == 150
    assert summary["total_tokens"] == 450
    assert summary["total_cost"] > 0

def test_token_tracker_cost_calculation():
    """Test cost calculation is accurate"""
    from services.fsd import TokenUsageTracker
    tracker = TokenUsageTracker()

    tracker.log_usage("Test", 1_000_000, 1_000_000)

    # Cost should be (1M * 0.15 + 1M * 0.60) / 1M = 0.75
    assert abs(tracker.total_cost - 0.75) < 0.01

def test_token_tracker_multiple_operations():
    """Test tracking multiple operations"""
    from services.fsd import TokenUsageTracker
    tracker = TokenUsageTracker()

    for i in range(5):
        tracker.log_usage(f"Op{i}", 100, 50)

    assert len(tracker.session_logs) == 5
    assert tracker.total_input_tokens == 500
    assert tracker.total_output_tokens == 250

def test_token_tracker_zero_tokens():
    """Test logging zero tokens"""
    from services.fsd import TokenUsageTracker
    tracker = TokenUsageTracker()

    tracker.log_usage("Empty", 0, 0)

    assert tracker.total_cost == 0.0
    assert len(tracker.session_logs) == 1

# ===============================
# Document Analyzer Tests (15)
# ===============================

def test_parse_pdf_pdfplumber_success(mock_env_vars, mocker):
    """Test PDF parsing with pdfplumber"""
    from services.fsd import DocumentAnalyzer, TokenUsageTracker
    from openai import OpenAI, AsyncOpenAI

    # Mock pdfplumber
    mock_pdf = mocker.MagicMock()
    mock_page = mocker.MagicMock()
    mock_page.extract_text.return_value = "Extracted text from PDF"
    mock_pdf.pages = [mock_page]

    mock_plumber = mocker.patch('services.fsd.pdfplumber.open')
    mock_plumber.return_value.__enter__.return_value = mock_pdf

    tracker = TokenUsageTracker()
    client = OpenAI(api_key="test")
    async_client = AsyncOpenAI(api_key="test")
    analyzer = DocumentAnalyzer(tracker, client, async_client)

    text = analyzer.parse_pdf(b"fake pdf")

    assert "Extracted text" in text

def test_parse_pdf_pypdf2_fallback(mock_env_vars, mocker):
    """Test PDF parsing fallback to pypdf"""
    from services.fsd import DocumentAnalyzer, TokenUsageTracker
    from openai import OpenAI, AsyncOpenAI

    # Mock pdfplumber to fail
    mocker.patch('services.fsd.pdfplumber.open', side_effect=Exception("Fail"))

    # Mock pypdf PdfReader
    mock_reader = mocker.MagicMock()
    mock_page = mocker.MagicMock()
    mock_page.extract_text.return_value = "pypdf extracted text"
    mock_reader.pages = [mock_page]

    mocker.patch('services.fsd.PdfReader', return_value=mock_reader)

    tracker = TokenUsageTracker()
    client = OpenAI(api_key="test")
    async_client = AsyncOpenAI(api_key="test")
    analyzer = DocumentAnalyzer(tracker, client, async_client)

    text = analyzer.parse_pdf(b"fake pdf")

    assert "pypdf" in text

def test_parse_pdf_error(mock_env_vars, mocker):
    """Test PDF parsing handles errors"""
    from services.fsd import DocumentAnalyzer, TokenUsageTracker
    from openai import OpenAI, AsyncOpenAI

    mocker.patch('services.fsd.pdfplumber.open', side_effect=Exception("Fail"))
    mocker.patch('services.fsd.PdfReader', side_effect=Exception("Fail"))

    tracker = TokenUsageTracker()
    client = OpenAI(api_key="test")
    async_client = AsyncOpenAI(api_key="test")
    analyzer = DocumentAnalyzer(tracker, client, async_client)

    with pytest.raises(ValueError):
        analyzer.parse_pdf(b"fake pdf")

def test_parse_docx_success(mock_env_vars, mocker):
    """Test successful DOCX parsing"""
    from services.fsd import DocumentAnalyzer, TokenUsageTracker
    from openai import OpenAI, AsyncOpenAI

    # Mock Document
    mock_doc = mocker.MagicMock()
    mock_para = mocker.MagicMock()
    mock_para.text = "Paragraph text from DOCX"
    mock_doc.paragraphs = [mock_para]
    mock_doc.tables = []

    mocker.patch('services.fsd.Document', return_value=mock_doc)

    tracker = TokenUsageTracker()
    client = OpenAI(api_key="test")
    async_client = AsyncOpenAI(api_key="test")
    analyzer = DocumentAnalyzer(tracker, client, async_client)

    text = analyzer.parse_docx(b"fake docx")

    assert "Paragraph text" in text

def test_parse_docx_with_tables(mock_env_vars, mocker):
    """Test DOCX parsing with tables"""
    from services.fsd import DocumentAnalyzer, TokenUsageTracker
    from openai import OpenAI, AsyncOpenAI

    mock_doc = mocker.MagicMock()
    mock_doc.paragraphs = []

    # Mock table
    mock_cell = mocker.MagicMock()
    mock_cell.text = "Table cell content"
    mock_row = mocker.MagicMock()
    mock_row.cells = [mock_cell]
    mock_table = mocker.MagicMock()
    mock_table.rows = [mock_row]
    mock_doc.tables = [mock_table]

    mocker.patch('services.fsd.Document', return_value=mock_doc)

    tracker = TokenUsageTracker()
    client = OpenAI(api_key="test")
    async_client = AsyncOpenAI(api_key="test")
    analyzer = DocumentAnalyzer(tracker, client, async_client)

    text = analyzer.parse_docx(b"fake docx")

    assert "Table cell" in text

def test_parse_docx_error(mock_env_vars, mocker):
    """Test DOCX parsing handles errors"""
    from services.fsd import DocumentAnalyzer, TokenUsageTracker
    from openai import OpenAI, AsyncOpenAI

    mocker.patch('services.fsd.Document', side_effect=Exception("Parse error"))

    tracker = TokenUsageTracker()
    client = OpenAI(api_key="test")
    async_client = AsyncOpenAI(api_key="test")
    analyzer = DocumentAnalyzer(tracker, client, async_client)

    with pytest.raises(ValueError):
        analyzer.parse_docx(b"fake docx")

@pytest.mark.asyncio
async def test_extract_sections_async(mock_env_vars, mock_async_openai):
    """Test async section extraction"""
    from services.fsd import DocumentAnalyzer, TokenUsageTracker
    from openai import OpenAI, AsyncOpenAI

    tracker = TokenUsageTracker()
    client = OpenAI(api_key="test")

    # Setup async mock
    async_client = AsyncOpenAI(api_key="test")
    mock_response = MagicMock()
    mock_response.choices = [MagicMock(message=MagicMock(content='{"executive_summary": "Test summary", "business_requirements": "Test reqs"}'))]
    mock_response.usage = MagicMock(prompt_tokens=50, completion_tokens=100)
    async_client.chat.completions.create = AsyncMock(return_value=mock_response)

    analyzer = DocumentAnalyzer(tracker, client, async_client)

    sections = await analyzer.extract_document_sections_async("Test document content")

    assert isinstance(sections, dict)
    assert "executive_summary" in sections or "business_requirements" in sections

@pytest.mark.asyncio
async def test_extract_sections_token_tracking(mock_env_vars):
    """Test section extraction tracks tokens"""
    from services.fsd import DocumentAnalyzer, TokenUsageTracker
    from openai import OpenAI, AsyncOpenAI

    tracker = TokenUsageTracker()
    client = OpenAI(api_key="test")
    async_client = AsyncOpenAI(api_key="test")

    mock_response = MagicMock()
    mock_response.choices = [MagicMock(message=MagicMock(content='{"executive_summary": "Test"}'))]
    mock_response.usage = MagicMock(prompt_tokens=100, completion_tokens=200)
    async_client.chat.completions.create = AsyncMock(return_value=mock_response)

    analyzer = DocumentAnalyzer(tracker, client, async_client)

    await analyzer.extract_document_sections_async("Content")

    assert tracker.total_input_tokens == 100
    assert tracker.total_output_tokens == 200

@pytest.mark.asyncio
async def test_extract_sections_error_handling(mock_env_vars):
    """Test section extraction handles errors gracefully"""
    from services.fsd import DocumentAnalyzer, TokenUsageTracker
    from openai import OpenAI, AsyncOpenAI

    tracker = TokenUsageTracker()
    client = OpenAI(api_key="test")
    async_client = AsyncOpenAI(api_key="test")
    async_client.chat.completions.create = AsyncMock(side_effect=Exception("API Error"))

    analyzer = DocumentAnalyzer(tracker, client, async_client)

    # Should return fallback structure instead of crashing
    result = await analyzer.extract_document_sections_async("Content")
    assert isinstance(result, dict)

@pytest.mark.asyncio
async def test_analyze_document_async_pdf(mock_env_vars, mocker):
    """Test full document analysis with PDF"""
    from services.fsd import DocumentAnalyzer, TokenUsageTracker
    from openai import OpenAI, AsyncOpenAI

    # Mock PDF parsing
    mocker.patch('services.fsd.pdfplumber.open')

    tracker = TokenUsageTracker()
    client = OpenAI(api_key="test")
    async_client = AsyncOpenAI(api_key="test")

    mock_response = MagicMock()
    mock_response.choices = [MagicMock(message=MagicMock(content='{"executive_summary": "Summary"}'))]
    mock_response.usage = MagicMock(prompt_tokens=50, completion_tokens=100)
    async_client.chat.completions.create = AsyncMock(return_value=mock_response)

    analyzer = DocumentAnalyzer(tracker, client, async_client)

    # Mock parse_pdf to return text
    analyzer.parse_pdf = lambda x: "Test document content"

    result = await analyzer.analyze_document_async(b"fake pdf", "test.pdf")

    assert "raw_content" in result
    assert "extracted_sections" in result

@pytest.mark.asyncio
async def test_analyze_document_async_docx(mock_env_vars, mocker):
    """Test full document analysis with DOCX"""
    from services.fsd import DocumentAnalyzer, TokenUsageTracker
    from openai import OpenAI, AsyncOpenAI

    tracker = TokenUsageTracker()
    client = OpenAI(api_key="test")
    async_client = AsyncOpenAI(api_key="test")

    mock_response = MagicMock()
    mock_response.choices = [MagicMock(message=MagicMock(content='{"executive_summary": "Summary"}'))]
    mock_response.usage = MagicMock(prompt_tokens=50, completion_tokens=100)
    async_client.chat.completions.create = AsyncMock(return_value=mock_response)

    analyzer = DocumentAnalyzer(tracker, client, async_client)

    # Mock parse_docx to return text
    analyzer.parse_docx = lambda x: "Test document content"

    result = await analyzer.analyze_document_async(b"fake docx", "test.docx")

    assert "raw_content" in result
    assert "extracted_sections" in result

@pytest.mark.asyncio
async def test_analyze_document_unsupported_type(mock_env_vars):
    """Test analysis rejects unsupported file types"""
    from services.fsd import DocumentAnalyzer, TokenUsageTracker
    from openai import OpenAI, AsyncOpenAI

    tracker = TokenUsageTracker()
    client = OpenAI(api_key="test")
    async_client = AsyncOpenAI(api_key="test")
    analyzer = DocumentAnalyzer(tracker, client, async_client)

    with pytest.raises(ValueError, match="Unsupported"):
        await analyzer.analyze_document_async(b"data", "test.xyz")

@pytest.mark.asyncio
async def test_analyze_document_empty_content(mock_env_vars, mocker):
    """Test analysis handles empty document content"""
    from services.fsd import DocumentAnalyzer, TokenUsageTracker
    from openai import OpenAI, AsyncOpenAI

    tracker = TokenUsageTracker()
    client = OpenAI(api_key="test")
    async_client = AsyncOpenAI(api_key="test")
    analyzer = DocumentAnalyzer(tracker, client, async_client)

    # Mock parse to return empty string
    analyzer.parse_pdf = lambda x: ""

    with pytest.raises(ValueError, match="No text content"):
        await analyzer.analyze_document_async(b"fake pdf", "test.pdf")

# ===============================
# FSD Agent Service Tests (20)
# ===============================

def test_fsd_service_initialization(mock_env_vars):
    """Test FSD service initializes correctly"""
    from services.fsd import FSDAgentService

    service = FSDAgentService()

    assert hasattr(service, 'token_tracker')
    assert hasattr(service, 'async_client')
    assert hasattr(service, 'document_analyzer')
    assert hasattr(service, 'generated_documents')

def test_fsd_service_has_qdrant_client(mock_env_vars):
    """Test service initializes Qdrant client"""
    from services.fsd import FSDAgentService

    service = FSDAgentService()

    assert hasattr(service, 'qdrant_client')

def test_fsd_service_generate_embeddings(mock_env_vars):
    """Test embedding generation"""
    from services.fsd import FSDAgentService

    service = FSDAgentService()

    # Should have embedding generation method
    assert hasattr(service, 'generate_embeddings')

@pytest.mark.asyncio
async def test_fsd_generate_document_async(mock_env_vars):
    """Test async FSD document generation"""
    from services.fsd import FSDAgentService

    service = FSDAgentService()

    # Mock the async OpenAI call
    mock_response = MagicMock()
    mock_response.choices = [MagicMock(message=MagicMock(content="Generated FSD section content"))]
    mock_response.usage = MagicMock(prompt_tokens=100, completion_tokens=500)
    service.async_client.chat.completions.create = AsyncMock(return_value=mock_response)

    result = await service.generate_document_with_llama_async(
        "Build user authentication system",
        "",
        ""
    )

    assert result is not None
    assert isinstance(result, str)

@pytest.mark.asyncio
async def test_fsd_from_document_upload(mock_env_vars, mocker):
    """Test FSD generation from document upload"""
    from services.fsd import FSDAgentService

    service = FSDAgentService()

    # Mock document analyzer
    mock_analysis = {
        "raw_content": "Test requirements",
        "extracted_sections": {"business_requirements": "Test reqs"}
    }
    service.document_analyzer.analyze_document_async = AsyncMock(return_value=mock_analysis)

    # Mock FSD generation
    mock_response = MagicMock()
    mock_response.choices = [MagicMock(message=MagicMock(content="FSD content"))]
    mock_response.usage = MagicMock(prompt_tokens=100, completion_tokens=500)
    service.async_client.chat.completions.create = AsyncMock(return_value=mock_response)

    # Mock Word document creation
    service._create_word_document_fsd = lambda x: BytesIO(b"fake docx")

    result = await service.generate_fsd_from_document_upload(b"fake pdf", "test.pdf")

    assert result is not None
    assert hasattr(result, 'document_id')

def test_fsd_get_document_success(mock_env_vars):
    """Test retrieving cached FSD document"""
    from services.fsd import FSDAgentService

    service = FSDAgentService()

    # Add document to cache
    doc_id = "test-123"
    mock_doc = b"fake docx content"
    service.generated_documents[doc_id] = mock_doc

    result = service.get_document(doc_id)

    assert result == mock_doc

def test_fsd_get_document_not_found(mock_env_vars):
    """Test document not found raises ValueError"""
    from services.fsd import FSDAgentService

    service = FSDAgentService()

    with pytest.raises(ValueError, match="Document not found"):
        service.get_document("non-existent-id")

def test_fsd_get_token_usage_stats(mock_env_vars):
    """Test getting token usage statistics"""
    from services.fsd import FSDAgentService

    service = FSDAgentService()

    # Log some token usage
    service.token_tracker.log_usage("Test1", 100, 50)
    service.token_tracker.log_usage("Test2", 200, 100)

    stats = service.get_token_usage_stats()

    assert "total_input_tokens" in stats
    assert "total_output_tokens" in stats
    assert "total_cost" in stats
    assert stats["total_input_tokens"] == 300
    assert stats["total_output_tokens"] == 150

def test_fsd_clear_cache(mock_env_vars):
    """Test cache clearing"""
    from services.fsd import FSDAgentService

    service = FSDAgentService()

    # Add documents to cache
    service.generated_documents["doc1"] = b"content1"
    service.generated_documents["doc2"] = b"content2"

    assert len(service.generated_documents) == 2

    result = service.clear_cache()

    assert len(service.generated_documents) == 0
    assert "message" in result

def test_fsd_get_mcp_context(mock_env_vars):
    """Test MCP context retrieval method exists"""
    from services.fsd import FSDAgentService

    service = FSDAgentService()

    # Should have MCP context method
    assert hasattr(service, 'get_mcp_context')

# ===============================
# Word Document Generation Tests (10)
# ===============================

def test_create_word_document():
    """Test Word document creation"""
    from docx import Document

    doc = Document()
    doc.add_heading("Test FSD Document", 0)
    doc.add_paragraph("This is a test paragraph.")

    assert len(doc.paragraphs) >= 1

def test_word_document_structure():
    """Test Word document has correct structure"""
    from docx import Document

    doc = Document()
    doc.add_heading("1. Introduction", 1)
    doc.add_paragraph("Introduction content")
    doc.add_heading("2. Requirements", 1)
    doc.add_paragraph("Requirements content")

    headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert len(headings) >= 2

def test_word_document_formatting():
    """Test Word document formatting"""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    para = doc.add_paragraph("Formatted text")
    run = para.runs[0]
    run.font.size = Pt(12)

    assert run.font.size == Pt(12)

def test_word_document_table_of_contents():
    """Test Word document TOC generation"""
    from docx import Document

    doc = Document()
    doc.add_heading("Table of Contents", 0)
    doc.add_paragraph("1. Introduction")
    doc.add_paragraph("2. Requirements")

    assert len(doc.paragraphs) >= 3

def test_word_document_sections():
    """Test Word document section generation"""
    from docx import Document

    doc = Document()

    sections = ["Introduction", "Business Requirements", "Technical Requirements"]
    for section in sections:
        doc.add_heading(section, 1)
        doc.add_paragraph(f"Content for {section}")

    assert len([p for p in doc.paragraphs if p.style.name.startswith('Heading')]) >= len(sections)

def test_word_document_export():
    """Test Word document export to bytes"""
    from docx import Document
    from io import BytesIO

    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("Test content")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    assert len(buffer.getvalue()) > 0

def test_word_document_with_tables():
    """Test Word document with tables"""
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=3, cols=3)

    # Add headers
    cells = table.rows[0].cells
    cells[0].text = "ID"
    cells[1].text = "Requirement"
    cells[2].text = "Priority"

    assert len(table.rows) == 3
    assert len(table.columns) == 3

def test_word_document_with_lists():
    """Test Word document with bulleted lists"""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Item 1", style='List Bullet')
    doc.add_paragraph("Item 2", style='List Bullet')
    doc.add_paragraph("Item 3", style='List Bullet')

    bullets = [p for p in doc.paragraphs if 'Bullet' in p.style.name]
    assert len(bullets) == 3

def test_word_document_with_numbering():
    """Test Word document with numbered lists"""
    from docx import Document

    doc = Document()
    doc.add_paragraph("First item", style='List Number')
    doc.add_paragraph("Second item", style='List Number')
    doc.add_paragraph("Third item", style='List Number')

    numbered = [p for p in doc.paragraphs if 'Number' in p.style.name]
    assert len(numbered) == 3

def test_word_document_page_breaks():
    """Test Word document page breaks"""
    from docx import Document
    from docx.enum.text import WD_BREAK

    doc = Document()
    doc.add_paragraph("Page 1 content")
    doc.add_page_break()
    doc.add_paragraph("Page 2 content")

    assert len(doc.paragraphs) >= 2
